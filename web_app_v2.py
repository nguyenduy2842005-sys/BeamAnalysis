"""
web_app_v2.py — Beam Analysis Suite (Đã sửa triệt để lỗi thụt dòng IndentationError)
Tabs:
  1. Single Beam     — giữ nguyên logic gốc
  2. Continuous Beam — FEM Euler-Bernoulli
  3. Plane Frame     — FEM 2D khung phẳng
"""

from __future__ import annotations

import io
import math
import os
import tempfile
import zipfile
from typing import Iterable
from xml.sax.saxutils import escape as _xml_escape

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import plotly.graph_objects as go
import streamlit as st
from st_aggrid import AgGrid, GridOptionsBuilder, GridUpdateMode, DataReturnMode

from beam_core import BeamInput, BeamResult, solve_beam
from fem_core import (
    SpanDef, SupportDef, ContinuousBeamInput, ContinuousBeamResult,
    solve_continuous_beam,
    FrameNode, FrameElement, FrameSupport, FramePointLoad,
    PlaneFrameInput, PlaneFrameResult, FrameElementResult,
    solve_plane_frame,
    build_docx_bytes,
)

# --- Hỗ trợ xuất ảnh chất lượng cao ---
try:
    import kaleido
    KALEIDO_AVAILABLE = True
except ImportError:
    KALEIDO_AVAILABLE = False
@st.cache_resource
def get_kaleido_chrome():
    try:
        kaleido.get_chrome_sync()
        return True
    except:
        return False

def ensure_kaleido_chrome() -> bool:
    return KALEIDO_AVAILABLE and get_kaleido_chrome()
def create_placeholder_image(text: str) -> bytes:
    """Tạo ảnh PNG với thông báo lỗi dùng Matplotlib."""
    import matplotlib.pyplot as plt
    import io
    fig, ax = plt.subplots(figsize=(8, 4.5))
    ax.text(0.5, 0.5, f"⚠️ Không thể hiển thị biểu đồ:\n{text}",
            ha='center', va='center', fontsize=14, transform=ax.transAxes)
    ax.set_axis_off()
    buf = io.BytesIO()
    plt.savefig(buf, format='png', dpi=200, bbox_inches='tight')
    plt.close()
    buf.seek(0)
    return buf.getvalue()
# ══════════════════════════════════════════════════════
#  GLOBAL CONSTANTS & SAFE ADAPTIVE THEME
# ══════════════════════════════════════════════════════
PLOT_HEIGHT = 315

COLOR_SFD   = "#0b5fff"
COLOR_BMD   = "#ff2b2b"
COLOR_ELAST = "#ff8800"
COLOR_AXIAL = "#9b27af"
COLOR_BEAM  = "#7a7f85"
COLOR_SUP   = "#ef1d14"

def _padded_range(values: np.ndarray, pad_frac: float = 0.18, min_span: float = 1.0) -> tuple[float, float]:
    """Tính khoảng [min,max] có đệm biên hợp lý cho trục y, dựa trên dữ liệu thật.
    Đảm bảo trục được KHÓA Ở MỘT TỶ LỆ CỐ ĐỊNH (không tự ý co giãn theo từng điểm hover),
    nhưng vẫn hiển thị trọn vẹn toàn bộ biểu đồ."""
    vmin = float(np.min(values))
    vmax = float(np.max(values))
    span = vmax - vmin
    if span < min_span:
        center = (vmax + vmin) / 2
        vmin, vmax = center - min_span / 2, center + min_span / 2
        span = vmax - vmin
    pad = span * pad_frac
    return (vmin - pad, vmax + pad)


def base_figure(title, length, y_title="", reverse_y=None):
    margin_x = max(length * 0.06, 0.5)

    fig = go.Figure()

    # Tự động nhận diện biểu đồ mô-men qua tiêu đề nếu không truyền tường minh reverse_y
    if reverse_y is None:
        title_lower = title.lower()
        reverse_y = any(kw in title_lower for kw in ["mô-men", "momen", "moment", "bmd"])

    # Range mặc định của bạn gốc là [-1.5, 1.2] (vmin = -1.5, vmax = 1.2)
    if reverse_y:
        y_range_config = [1.2, -1.5]  # [vmax, vmin] -> Ép trục Y đảo ngược cố định (Dương hướng xuống)
    else:
        y_range_config = [-1.5, 1.2]  # [vmin, vmax] -> Trục Y xuôi thông thường (Dương hướng lên)

    fig.update_layout(
        title=dict(
            text=f"<b>{title}</b>",
            x=0.5,
            xanchor="center",
            font=dict(size=14)
        ),
        height=315,
        margin=dict(l=55, r=20, t=60, b=75),
        paper_bgcolor="rgba(0,0,0,0)",
        plot_bgcolor="rgba(0,0,0,0)",
        showlegend=False,
        xaxis=dict(
            title="x (m)",
            range=[-margin_x, length + margin_x],
            zeroline=True,
            mirror=True,
            showgrid=True,
            linecolor="gray",
            gridcolor="rgba(128,128,128,0.2)"
        ),
        yaxis=dict(
            title=y_title,
            range=y_range_config,  # <--- Áp dụng cấu hình range đã chuẩn hóa hình học
            zeroline=True,
            mirror=True,
            showgrid=True,
            linecolor="gray",
            gridcolor="rgba(128,128,128,0.2)"
        )
    )
    return fig


def synced_figure(title, length, y_range=(-1.5, 1.2), y_title="", reverse_y=None):
    margin_x = max(length * 0.06, 0.5)

    fig = go.Figure()

    # Tự động nhận diện biểu đồ mô-men qua tiêu đề nếu không truyền tường minh reverse_y
    if reverse_y is None:
        title_lower = title.lower()
        reverse_y = any(kw in title_lower for kw in ["mô-men", "momen", "moment", "bmd"])

    # Xử lý dải giá trị y_range dựa trên việc có đảo trục hay không
    if y_range is not None:
        vmin = min(y_range)
        vmax = max(y_range)
        if reverse_y:
            actual_y_range = [vmax, vmin]  # Số lớn xếp trước số nhỏ -> Ép cả Web lẫn Kaleido lật ngược trục Y 100%
        else:
            actual_y_range = [vmin, vmax]  # Số nhỏ xếp trước số lớn -> Trục xuôi toán học thông thường
    else:
        actual_y_range = None

    fig.update_layout(
        title=dict(
            text=f"<b>{title}</b>",
            x=0.5,
            xanchor="center",
            font=dict(size=15)
        ),
        height=PLOT_HEIGHT,
        margin=dict(l=55, r=25, t=60, b=65),
        paper_bgcolor="rgba(0,0,0,0)",
        plot_bgcolor="rgba(0,0,0,0)",
        showlegend=False,
        xaxis=dict(
            title="x (m)",
            range=[-margin_x, length + margin_x],
            zeroline=True,
            mirror=True,
            scaleanchor=None
        ),
        yaxis=dict(
            title=y_title,
            range=actual_y_range,  # <--- Sử dụng range cấu hình đảo tường minh ở đây
            zeroline=True,
            mirror=True
        )
    )

    fig.update_xaxes(
        showgrid=True,
        linecolor="gray",
        gridcolor="rgba(128,128,128,0.2)"
    )
    fig.update_yaxes(
        showgrid=True,
        linecolor="gray",
        gridcolor="rgba(128,128,128,0.2)"
    )
    return fig

def single_synced_plot(title, length, y_title=""):
        return synced_figure(
            title,
            length,
            y_range=(-1.5, 1.2),
            y_title=y_title
        )

# ══════════════════════════════════════════════════════
#  PAGE CONFIG & ADAPTIVE CSS
# ══════════════════════════════════════════════════════
st.set_page_config(
    page_title="Beam Analysis ",
    page_icon="🏗️",
    layout="wide",
    initial_sidebar_state="expanded"
)

def inject_css() -> None:
    # Xác định trạng thái sidebar từ session_state
    sidebar_open = st.session_state.get("sidebar_open", True)

    # CSS ẩn/hiện sidebar dựa trên trạng thái
    sidebar_css = """
    section[data-testid="stSidebar"] {
        display: block !important;
        width: 21rem !important;
        min-width: 21rem !important;
        transition: all 0.25s ease;
    }
    """ if sidebar_open else """
    section[data-testid="stSidebar"] {
        display: none !important;
        width: 0px !important;
        min-width: 0px !important;
    }
    """

    st.markdown(f"""
    <style>
    /* Ẩn các thành phần không cần thiết */
    #MainMenu, footer, [data-testid="stToolbar"] {{
        display: none !important;
    }}

    /* Sidebar toggle CSS */
    {sidebar_css}

    /* Layout */
    .block-container {{
        padding-top: 0.5rem;
        padding-bottom: 1.5rem;
        max-width: 1680px;
    }}
    .metric-strip {{
        display: grid;
        grid-template-columns: repeat(4, minmax(0, 1fr));
        gap: 10px;
        margin: 0.2rem 0 0.8rem;
    }}
    .metric-card {{
        background: var(--background-color);
        border: 1px solid var(--secondary-background-color);
        border-radius: 6px;
        padding: 10px 12px;
        box-shadow: 0 1px 3px rgba(0,0,0,0.05);
    }}
    .metric-label {{
        color: var(--text-color);
        opacity: 0.7;
        font-size: 0.78rem;
        margin-bottom: 3px;
    }}
    .metric-value {{
        color: var(--text-color);
        font-weight: 700;
        font-size: 1.08rem;
    }}
    div[data-testid="stVerticalBlockBorderWrapper"] {{
        border-color: var(--secondary-background-color);
    }}
    .stPlotlyChart {{
        border: 1px solid var(--secondary-background-color);
        border-radius: 6px;
        padding: 6px;
        background-color: transparent !important;
    }}
    textarea {{
        font-family: Consolas, "Courier New", monospace !important;
    }}

    /* Style nút hamburger trong header */
    div[data-testid="stHorizontalBlock"] > div:first-child button {{
        padding: 4px 10px;
        font-size: 20px;
        line-height: 1;
    }}
    </style>
    """, unsafe_allow_html=True)

# ══════════════════════════════════════════════════════
#  SHARED HELPERS
# ══════════════════════════════════════════════════════

def clean_rows(df: pd.DataFrame, columns: Iterable[str]) -> list[tuple[float, ...]]:
    rows: list[tuple[float, ...]] = []
    for _, _row in df.iterrows():
        values: list[float] = []
        skip = False
        for col in columns:
            v = _row.get(col)
            if v is None or pd.isna(v) or v == "":
                skip = True; break
            try:
                n = float(v)
            except (TypeError, ValueError):
                skip = True; break
            if not math.isfinite(n):
                skip = True; break
            values.append(n)
        if not skip:
            rows.append(tuple(values))
    return rows


def safe_data_editor(
    widget_key: str,
    default_df: pd.DataFrame,
    **editor_kwargs
) -> pd.DataFrame:

    data_key = f"{widget_key}__data"

    if data_key not in st.session_state:
        st.session_state[data_key] = default_df.copy()

    edited_df = st.data_editor(
        st.session_state[data_key],
        key=widget_key,
        **editor_kwargs,
    )

    st.session_state[data_key] = edited_df.copy()

    return edited_df
    # Hiển thị Data Editor
    edited_df = st.data_editor(
        st.session_state[seed_key],
        key=widget_key,
        **editor_kwargs,
    )

    # Cập nhật Session State
    st.session_state[seed_key] = edited_df.copy()

    errors = []

    # Kiểm tra từng dòng
    if "x_local (m)" in edited_df.columns:

        for row_id, row in edited_df.iterrows():

            if pd.isna(row["x_local (m)"]):
                continue

            try:
                x = float(row["x_local (m)"])
            except Exception:
                continue

            if not (0.0 <= x <= L_span):

                errors.append(
                    f"• Dòng {row_id + 1}: "
                    f"x_local = {x:.2f} m "
                    f"(Giới hạn: 0.00 → {L_span:.2f} m)"
                )

    # Hiển thị lỗi
    if errors:

        st.error(
            f"❌ Vị trí tải trọng phải nằm trong khoảng "
            f"0.00 ≤ x_local ≤ {L_span:.2f} m"
        )

        for msg in errors:
            st.warning(msg)

    return edited_df, len(errors) > 0
    base = st.session_state[seed_key].copy()
    if isinstance(edited, dict):
        for ridx, changes in edited.get("edited_rows", {}).items():
            for col, val in changes.items():
                base.loc[int(ridx), col] = val
        for new_row in edited.get("added_rows", []):
            base = pd.concat([base, pd.DataFrame([new_row])], ignore_index=True)
        deleted = sorted(edited.get("deleted_rows", []), reverse=True)
        if deleted:
            base = base.drop(index=deleted).reset_index(drop=True)
    return base
def validate_local_coordinate(
    df: pd.DataFrame,
    L_span: float,
    x_column: str = "x_local (m)",
):
    """
    Kiểm tra x_local của Point Load hoặc Point Moment.

    Returns
    -------
    errors : list[str]
    """

    errors = []

    if x_column not in df.columns:
        return errors

    for idx, row in df.iterrows():

        value = row[x_column]

        if pd.isna(value):
            continue

        try:
            x = float(value)

        except Exception:

            errors.append(
                f"Dòng {idx+1}: {x_column} không phải số."
            )
            continue

        if x < 0:

            errors.append(
                f"Dòng {idx+1}: {x_column} = {x:.2f} m < 0."
            )

        elif x > L_span:

            errors.append(
                f"Dòng {idx+1}: {x_column} = {x:.2f} m > {L_span:.2f} m."
            )

    return errors
def validate_local_coordinate(
    df: pd.DataFrame,
    L_span: float,
    column: str = "x_local (m)",
) -> list[str]:
    """
    Kiểm tra tọa độ local của Point Load / Point Moment.
    """

    errors = []

    if column not in df.columns:
        return errors

    for i, row in df.iterrows():

        value = row[column]

        if value is None or value == "" or pd.isna(value):
            continue

        try:
            x = float(value)
        except (TypeError, ValueError):

            errors.append(
                f"Dòng {i+1}: {column} không phải số."
            )
            continue

        if x < 0:

            errors.append(
                f"Dòng {i+1}: {column} = {x:.2f} m < 0."
            )

        elif x > L_span:

            errors.append(
                f"Dòng {i+1}: {column} = {x:.2f} m > {L_span:.2f} m."
            )

    return errors
def reset_keys_with_prefix(*prefixes: str) -> None:
    """Xóa tất cả session_state key theo prefix VÀ tăng reset_counter
    để buộc các widget AgGrid (và mọi widget có key động) phải
    render lại từ đầu với data rỗng khi New Model được bấm."""
    for k in list(st.session_state.keys()):
        if any(k.startswith(p) for p in prefixes):
            st.session_state.pop(k, None)
    # Tăng counter riêng cho mỗi prefix → key widget AgGrid thay đổi →
    # Streamlit tạo widget mới → data trong browser DOM bị clear hoàn toàn
    for p in prefixes:
        counter_key = f"__reset_cnt_{p}"
        st.session_state[counter_key] = st.session_state.get(counter_key, 0) + 1


def _get_reset_cnt(prefix: str) -> int:
    """Trả về reset counter hiện tại cho prefix (mặc định 0)."""
    return st.session_state.get(f"__reset_cnt_{prefix}", 0)

def validate_point_loads(rows, L):
    errors = []

    for i, row in enumerate(rows, start=1):
        try:
            P, x = row
        except Exception:
            continue

        if x < 0 or x > L:
            errors.append(
                f"Point Load - Dòng {i}: "
                f"x_local = {x:.3f} m, "
                f"phải nằm trong [0 ; {L:.3f}] m."
            )

    return errors


def validate_point_moments(rows, L):
    errors = []

    for i, row in enumerate(rows, start=1):
        try:
            M, x = row
        except Exception:
            continue

        if x < 0 or x > L:
            errors.append(
                f"Point Moment - Dòng {i}: "
                f"x_local = {x:.3f} m, "
                f"phải nằm trong [0 ; {L:.3f}] m."
            )

    return errors


def validate_udls(rows, L):
    errors = []

    for i, row in enumerate(rows, start=1):
        try:
            q, x1, x2 = row
        except Exception:
            continue

        if x1 < 0:
            errors.append(
                f"UDL - Dòng {i}: x1_local < 0."
            )

        if x2 > L:
            errors.append(
                f"UDL - Dòng {i}: x2_local > {L:.3f} m."
            )

        if x2 <= x1:
            errors.append(
                f"UDL - Dòng {i}: x2_local phải lớn hơn x1_local."
            )

    return errors
def metric_html(values: list[tuple[str, str]]) -> None:
    cards = "".join(
        f"<div class='metric-card'>"
        f"<div class='metric-label'>{lbl}</div>"
        f"<div class='metric-value'>{val}</div></div>"
        for lbl, val in values
    )
    st.markdown(f"<div class='metric-strip'>{cards}</div>", unsafe_allow_html=True)


def _minimal_docx_bytes(report_text: str, title: str = "Thuyết Minh Tính Toán") -> bytes:
    """
    Tạo file .docx hợp lệ TỐI THIỂU bằng cách ghi trực tiếp định dạng OOXML
    (ZIP + XML chuẩn của Word), KHÔNG phụ thuộc thư viện python-docx.

    Đây là lưới an toàn cuối cùng: nếu vì lý do nào đó server chưa cài được
    python-docx (ví dụ quên cập nhật requirements.txt), người dùng web VẪN
    luôn tải được file .docx hợp lệ — không bao giờ thấy lỗi "hãy tự cài thư viện".
    """


    def esc(s: str) -> str:
        return _xml_escape(s).replace("\t", "    ")

    body_paragraphs = []
    title_xml = (
        f'<w:p><w:pPr><w:jc w:val="center"/></w:pPr>'
        f'<w:r><w:rPr><w:b/><w:sz w:val="32"/></w:rPr>'
        f'<w:t xml:space="preserve">{esc(title)}</w:t></w:r></w:p>'
    )
    body_paragraphs.append(title_xml)
    body_paragraphs.append('<w:p/>')

    for line in report_text.split("\n"):
        stripped = line.strip()
        if stripped.startswith("="):
            continue
        is_heading = bool(stripped) and stripped[0].isdigit() and ". " in stripped[:4]
        if is_heading:
            run = (f'<w:r><w:rPr><w:b/><w:sz w:val="24"/></w:rPr>'
                   f'<w:t xml:space="preserve">{esc(line)}</w:t></w:r>')
        else:
            run = (f'<w:r><w:rPr><w:rFonts w:ascii="Consolas" w:hAnsi="Consolas"/>'
                   f'<w:sz w:val="18"/></w:rPr>'
                   f'<w:t xml:space="preserve">{esc(line) if line.strip() else " "}</w:t></w:r>')
        body_paragraphs.append(f'<w:p>{run}</w:p>')

    body_xml = "".join(body_paragraphs)

    document_xml = f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
<w:body>{body_xml}
<w:sectPr><w:pgSz w:w="11906" w:h="16838"/>
<w:pgMar w:top="1134" w:right="1134" w:bottom="1134" w:left="1417"/></w:sectPr>
</w:body></w:document>'''

    content_types = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
<Default Extension="xml" ContentType="application/xml"/>
<Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
</Types>'''

    rels = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>
</Relationships>'''

    doc_rels = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
</Relationships>'''

    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as z:
        z.writestr("[Content_Types].xml", content_types)
        z.writestr("_rels/.rels", rels)
        z.writestr("word/document.xml", document_xml)
        z.writestr("word/_rels/document.xml.rels", doc_rels)
    return buf.getvalue()

def plotly_to_png_fallback(fig: go.Figure) -> bytes | None:
    """
    Fallback dùng Matplotlib, vẽ lại trace, marker, shape và annotation (text, mũi tên đúng chiều).
    Không cần Chrome/kaleido.
    """
    try:
        import matplotlib.pyplot as plt
        import matplotlib.patches as patches
        import io
        import numpy as np

        plt.figure(figsize=(8, 4.5))

        # 1. Vẽ các trace (lines, markers, fills)
        for tr in fig.data:
            if hasattr(tr, 'x') and hasattr(tr, 'y') and tr.x is not None and tr.y is not None:
                try:
                    x = np.array(tr.x, dtype=float)
                    y = np.array(tr.y, dtype=float)
                    color = tr.line.color if hasattr(tr, 'line') and tr.line else '#0000ff'
                    linewidth = tr.line.width if hasattr(tr, 'line') and tr.line else 1.5

                    mode = tr.mode if hasattr(tr, 'mode') else 'lines'
                    if 'lines' in mode:
                        plt.plot(x, y, color=color, linewidth=linewidth)

                    if 'markers' in mode:
                        marker_size = tr.marker.size if hasattr(tr, 'marker') and tr.marker else 7
                        marker_color = tr.marker.color if hasattr(tr, 'marker') and tr.marker else color
                        plt.scatter(x, y, s=marker_size**2, color=marker_color, edgecolors='black', linewidth=0.5)

                    if hasattr(tr, 'fill') and tr.fill == 'tozeroy':
                        plt.fill_between(x, y, 0, color=color, alpha=0.2)
                    elif hasattr(tr, 'fill') and tr.fill == 'toself':
                        plt.fill(x, y, color=color, alpha=0.2)
                except Exception:
                    pass

        # 2. Vẽ các shape (ngàm, tam giác gối, v.v.)
        if fig.layout.shapes:
            for sh in fig.layout.shapes:
                try:
                    if sh.type == 'rect':
                        x0 = sh.x0 if sh.x0 is not None else 0
                        x1 = sh.x1 if sh.x1 is not None else 0
                        y0 = sh.y0 if sh.y0 is not None else 0
                        y1 = sh.y1 if sh.y1 is not None else 0
                        rect = patches.Rectangle(
                            (x0, y0), x1 - x0, y1 - y0,
                            facecolor=sh.fillcolor if sh.fillcolor else 'gray',
                            edgecolor=sh.line.color if sh.line else 'black',
                            linewidth=sh.line.width if sh.line else 1
                        )
                        plt.gca().add_patch(rect)
                except Exception:
                    pass

        # 3. Vẽ các annotation (text, mũi tên) – chỉ khi annotations không rỗng
        if fig.layout.annotations:
            for ann in fig.layout.annotations:
                try:
                    # Text (không có mũi tên)
                    if ann.text and not ann.showarrow:
                        x = ann.x if ann.x is not None else 0
                        y = ann.y if ann.y is not None else 0
                        plt.text(x, y, ann.text,
                                 fontsize=ann.font.size if ann.font else 10,
                                 color=ann.font.color if ann.font else 'black',
                                 ha=ann.xanchor if ann.xanchor else 'center',
                                 va=ann.yanchor if ann.yanchor else 'center')

                    # Mũi tên: vẽ từ (x, y) đến (ax, ay) (đúng với Plotly)
                    if ann.showarrow:
                        # Tail là (ann.ax, ann.ay), Head là (ann.x, ann.y)
                        x_tail = ann.ax if ann.ax is not None else 0
                        y_tail = ann.ay if ann.ay is not None else 0
                        x_head = ann.x if ann.x is not None else 0
                        y_head = ann.y if ann.y is not None else 0

                        plt.annotate('', xy=(x_head, y_head), xytext=(x_tail, y_tail),
                                     arrowprops=dict(arrowstyle='->',
                                                     color=ann.arrowcolor if ann.arrowcolor else 'black',
                                                     lw=ann.arrowwidth if ann.arrowwidth else 1.5))
                except Exception:
                    pass

        # 4. Tiêu đề, nhãn trục, grid
        title = fig.layout.title.text if fig.layout.title else ''
        plt.title(title)
        if fig.layout.xaxis and fig.layout.xaxis.title:
            plt.xlabel(fig.layout.xaxis.title.text)
        if fig.layout.yaxis and fig.layout.yaxis.title:
            plt.ylabel(fig.layout.yaxis.title.text)
        plt.grid(True)

        # Đồng bộ khoảng trục
        if fig.layout.xaxis and fig.layout.xaxis.range:
            plt.xlim(fig.layout.xaxis.range)
        if fig.layout.yaxis and fig.layout.yaxis.range:
            plt.ylim(fig.layout.yaxis.range)

        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=200, bbox_inches='tight')
        plt.close()
        buf.seek(0)
        return buf.getvalue()
    except Exception:
        return None
def docx_with_images(
    report_text: str,
    report_title: str,
    figures: list[tuple[str, go.Figure]]
) -> bytes:
    """
    Tạo file DOCX với văn bản và các biểu đồ.
    Ưu tiên dùng kaleido để có ảnh đẹp, nếu không được thì fallback sang Matplotlib.
    """
    from docx import Document
    from docx.shared import Inches
    import tempfile
    import os

    doc = Document()
    doc.add_heading(report_title, level=1)

    for line in report_text.split('\n'):
        doc.add_paragraph(line)

    if figures:
        doc.add_page_break()
        doc.add_heading('Biểu đồ kết quả', level=1)

        # Chuẩn bị kaleido nếu có thể
        use_kaleido = ensure_kaleido_chrome()

        for name, fig in figures:
            doc.add_heading(name, level=2)

            img_bytes = None

            # 1. Thử dùng kaleido (chất lượng cao)
            if use_kaleido:
                try:
                    img_bytes = fig.to_image(
                        format='png',
                        width=1600,
                        height=900,
                        scale=2,
                        engine='kaleido'
                    )
                except Exception:
                    img_bytes = None

            # 2. Nếu thất bại, dùng fallback Matplotlib
            if img_bytes is None:
                img_bytes = plotly_to_png_fallback(fig)
                if img_bytes is None:
                    # 3. Thậm chí fallback cũng lỗi → placeholder
                    img_bytes = create_placeholder_image(name)

            # Chèn ảnh vào Word
            with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp:
                tmp_path = tmp.name
            try:
                with open(tmp_path, 'wb') as f:
                    f.write(img_bytes)
                doc.add_picture(tmp_path, width=Inches(6.5))
            finally:
                if os.path.exists(tmp_path):
                    os.unlink(tmp_path)

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()
def docx_with_images(
    report_text: str,
    report_title: str,
    figures: list[tuple[str, go.Figure]]
) -> bytes:
    """Create a polished engineering calculation DOCX with sharp figures."""
    from docx import Document
    from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Inches, Pt, RGBColor
    import os
    import re
    import tempfile

    BLUE = RGBColor(31, 77, 120)
    ACCENT = RGBColor(46, 116, 181)
    MUTED = RGBColor(95, 105, 115)
    LIGHT_FILL = "F2F4F7"
    OK_FILL = "EAF4ED"
    CONTENT_WIDTH_DXA = 9360

    def set_run_font(run, name="Calibri", size=11, color=None, bold=None, italic=None):
        run.font.name = name
        run._element.rPr.rFonts.set(qn("w:ascii"), name)
        run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
        run._element.rPr.rFonts.set(qn("w:cs"), name)
        run.font.size = Pt(size)
        if color is not None:
            run.font.color.rgb = color
        if bold is not None:
            run.bold = bold
        if italic is not None:
            run.italic = italic

    def set_cell_fill(cell, fill):
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = tc_pr.find(qn("w:shd"))
        if shd is None:
            shd = OxmlElement("w:shd")
            tc_pr.append(shd)
        shd.set(qn("w:fill"), fill)

    def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_mar = tc_pr.first_child_found_in("w:tcMar")
        if tc_mar is None:
            tc_mar = OxmlElement("w:tcMar")
            tc_pr.append(tc_mar)
        for key, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
            node = tc_mar.find(qn(f"w:{key}"))
            if node is None:
                node = OxmlElement(f"w:{key}")
                tc_mar.append(node)
            node.set(qn("w:w"), str(value))
            node.set(qn("w:type"), "dxa")

    def set_table_width(table, widths_dxa):
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        tbl_pr = table._tbl.tblPr
        tbl_w = tbl_pr.find(qn("w:tblW"))
        if tbl_w is None:
            tbl_w = OxmlElement("w:tblW")
            tbl_pr.append(tbl_w)
        tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
        tbl_w.set(qn("w:type"), "dxa")
        grid = table._tbl.tblGrid
        if grid is None:
            grid = OxmlElement("w:tblGrid")
            table._tbl.insert(0, grid)
        for child in list(grid):
            grid.remove(child)
        for width in widths_dxa:
            col = OxmlElement("w:gridCol")
            col.set(qn("w:w"), str(width))
            grid.append(col)
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                width = widths_dxa[min(i, len(widths_dxa) - 1)]
                cell.width = Pt(width / 20)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                set_cell_margins(cell)
                tc_pr = cell._tc.get_or_add_tcPr()
                tc_w = tc_pr.find(qn("w:tcW"))
                if tc_w is None:
                    tc_w = OxmlElement("w:tcW")
                    tc_pr.append(tc_w)
                tc_w.set(qn("w:w"), str(width))
                tc_w.set(qn("w:type"), "dxa")

    def add_rule(paragraph, color="2E74B5", size="10", space="8"):
        p_pr = paragraph._p.get_or_add_pPr()
        p_bdr = p_pr.find(qn("w:pBdr"))
        if p_bdr is None:
            p_bdr = OxmlElement("w:pBdr")
            p_pr.append(p_bdr)
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), size)
        bottom.set(qn("w:space"), space)
        bottom.set(qn("w:color"), color)
        p_bdr.append(bottom)

    def add_kv_table(doc, rows, widths=(2300, 7060)):
        table = doc.add_table(rows=0, cols=2)
        table.style = "Table Grid"
        for label, value in rows:
            cells = table.add_row().cells
            cells[0].text = str(label)
            cells[1].text = str(value)
            set_cell_fill(cells[0], LIGHT_FILL)
        set_table_width(table, list(widths))
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                for p in cell.paragraphs:
                    p.paragraph_format.space_after = Pt(0)
                    for r in p.runs:
                        set_run_font(r, size=9.5, color=BLUE if i == 0 else RGBColor(30, 30, 30), bold=(i == 0))
        return table

    def add_callout(doc, text, fill=OK_FILL):
        table = doc.add_table(rows=1, cols=1)
        table.style = "Table Grid"
        cell = table.cell(0, 0)
        cell.text = ""
        set_cell_fill(cell, fill)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        set_run_font(r, size=10.5, color=RGBColor(28, 88, 48), bold=True)
        set_table_width(table, [CONTENT_WIDTH_DXA])

    def add_report_line(doc, line):
        stripped = line.strip()
        if not stripped or set(stripped) <= {"=", "-", " "}:
            return
        if re.match(r"^\d+\.\s+", stripped):
            doc.add_heading(stripped, level=1)
            return
        if re.match(r"^[a-z]\.\s+", stripped, flags=re.I):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(3)
            r = p.add_run(stripped)
            set_run_font(r, size=10.5, color=BLUE, bold=True)
            return
        looks_tabular = "|" in line or "  " in line or re.search(r"[+-]?\d+\.\d{3,}", line)
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(line)
        set_run_font(
            r,
            name="Consolas" if looks_tabular else "Calibri",
            size=8.8 if looks_tabular else 10.5,
            color=RGBColor(35, 35, 35),
        )

    def clean_title(text):
        return re.sub(r"<[^>]+>", "", str(text or "")).strip()

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.59)
    section.page_height = Cm(27.94)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for style_name, size, color, before, after in [
        ("Heading 1", 16, ACCENT, 16, 8),
        ("Heading 2", 13, ACCENT, 12, 6),
        ("Heading 3", 12, BLUE, 8, 4),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header_p = section.header.paragraphs[0]
    header_p.text = "Thuyết minh tính toán kết cấu | Beam Analysis Toolbox"
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for r in header_p.runs:
        set_run_font(r, size=8.5, color=MUTED)
    footer_p = section.footer.paragraphs[0]
    footer_p.text = "Tài liệu tính toán tự động - cần được kỹ sư phụ trách kiểm tra và phê duyệt trước khi phát hành."
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in footer_p.runs:
        set_run_font(r, size=8, color=MUTED)

    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_after = Pt(4)
    r = title_p.add_run("THUYẾT MINH TÍNH TOÁN")
    set_run_font(r, size=24, color=BLUE, bold=True)
    subtitle_p = doc.add_paragraph()
    subtitle_p.paragraph_format.space_after = Pt(12)
    r = subtitle_p.add_run(report_title)
    set_run_font(r, size=13, color=MUTED, bold=True)
    add_rule(subtitle_p)

    add_kv_table(doc, [
        ("Phần mềm", "Beam Analysis Toolbox"),
        ("Hạng mục", report_title),
        ("Phương pháp chính", "Phần tử hữu hạn khung phẳng 2D"),
        ("Phương pháp đối chiếu", "Maxwell-Mohr / nhân biểu đồ Vereshchagin"),
        ("Đơn vị", "kN, m, kNm"),
    ])
    doc.add_paragraph()

    max_diff = re.search(r"Sai khác lớn nhất:\s*([^\n]+)", report_text)
    if max_diff:
        add_callout(doc, f"Kiểm tra độc lập FEM - Vereshchagin: sai khác lớn nhất {max_diff.group(1).strip()}.")
    else:
        add_callout(doc, "Thuyết minh bao gồm thiết lập mô hình, điều kiện biên, nội lực, phản lực và kiểm tra cân bằng.")

    doc.add_page_break()
    doc.add_heading("Nội dung thuyết minh", level=1)
    for line in report_text.split("\n"):
        add_report_line(doc, line)

    if figures:
        doc.add_page_break()
        doc.add_heading("Biểu đồ kết quả", level=1)
        use_kaleido = ensure_kaleido_chrome()
        for idx, (name, fig) in enumerate(figures, 1):
            img_bytes = None
            if use_kaleido:
                try:
                    img_bytes = fig.to_image(format="png", width=1800, height=1050, scale=2, engine="kaleido")
                except Exception:
                    img_bytes = None
            if img_bytes is None:
                img_bytes = plotly_to_png_fallback(fig) or create_placeholder_image(name)

            with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
                tmp_path = tmp.name
            try:
                with open(tmp_path, "wb") as f:
                    f.write(img_bytes)
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.keep_with_next = True
                p.add_run().add_picture(tmp_path, width=Inches(6.35))
                cap = doc.add_paragraph()
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap.paragraph_format.space_after = Pt(10)
                r = cap.add_run(f"Hình {idx}. {clean_title(name)}")
                set_run_font(r, size=9.5, color=MUTED, italic=True)
            finally:
                if os.path.exists(tmp_path):
                    os.unlink(tmp_path)

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def report_panel(
    report_text: str | None,
    report_title: str,
    key_prefix: str,
    figures=None

) -> None:
    st.subheader("📋 Thuyết minh tính toán")
    if not report_text:
        st.info("Chưa có kết quả. Nhấn **▶ Solve** để tính toán và xem thuyết minh.")
        return

    st.code(report_text, language=None, line_numbers=False)
    st.markdown("---")
    st.caption("✅ Đọc kỹ thuyết minh ở trên. Nếu kết quả hợp lý, bạn có thể xuất file bên dưới.")

    c1, c2 = st.columns(2)
    with c1:
        st.download_button(
            "⬇️ Xuất .txt",
            data=report_text.encode("utf-8"),
            file_name=f"{key_prefix}_report.txt",
            mime="text/plain",
            use_container_width=True,
            key=f"{key_prefix}_dl_txt",
        )
    with c2:

        try:

            if figures:
                docx_bytes = docx_with_images(
                    report_text,
                    report_title,
                    figures
                )
            else:
                docx_bytes = build_docx_bytes(
                    report_text,
                    report_title
                )

        except ImportError:

            # fallback không cần python-docx
            docx_bytes = _minimal_docx_bytes(
                report_text,
                report_title
            )

        except Exception as e:

            st.warning(
                f"DOCX nâng cao lỗi ({e}). "
                "Đang dùng DOCX tối giản."
            )

            docx_bytes = _minimal_docx_bytes(
                report_text,
                report_title
            )

        st.download_button(
            "⬇️ Xuất .docx",
            data=docx_bytes,
            file_name=f"{key_prefix}_report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
            key=f"{key_prefix}_dl_docx",
        )


# ══════════════════════════════════════════════════════
#  ── TAB 1: SINGLE BEAM ──────────────────────────────
# ══════════════════════════════════════════════════════

def validate_single(data: BeamInput) -> list[str]:
    errors: list[str] = []
    if data.length <= 0:
        errors.append("Chiều dài dầm phải lớn hơn 0.")
    for name, rows in [("Point Load", [(x,) for _, x in data.point_loads]),
                        ("Point Moment", [(x,) for _, x in data.point_moments])]:
        for i, (xp,) in enumerate(rows, 1):
            if xp < 0 or xp > data.length:
                errors.append(f"{name} dòng {i}: vị trí x phải nằm trong [0, L].")
    for name, rows in [("UDL", data.udls), ("UVL", data.uvls)]:
        for i, (_, x1, x2) in enumerate(rows, 1):
            if x1 < 0 or x2 < 0 or x1 > data.length or x2 > data.length or x2 <= x1:
                errors.append(f"{name} dòng {i}: cần 0 ≤ x1 < x2 ≤ L.")
    return errors


def draw_supports_single(fig: go.Figure, data: BeamInput) -> None:
    l = data.length
    if data.beam_type == "simple":
        # Node 0 (Gối cố định - Pin: giữ nguyên hình tam giác)
        fig.add_trace(go.Scatter(
            x=[0, l / 34, -l / 34, 0], y=[0, -0.37, -0.37, 0],
            fill="toself", mode="lines",
            line={"color": COLOR_SUP, "width": 1.5}, fillcolor=COLOR_SUP,
            hoverinfo="skip"))

        # Node L (Gối di động dầm đơn - Tinh chỉnh khoảng cách thoáng, đẹp)
        y_top = -0.08
        y_bot = -0.36
        y_floor = -0.44

        # Điểm chấm 1 (Phía trên)
        fig.add_trace(go.Scatter(
            x=[l], y=[y_top], mode="markers",
            marker=dict(symbol="circle", size=7, color=COLOR_SUP, line=dict(width=1, color=COLOR_SUP)),
            hoverinfo="skip"
        ))
        # Điểm chấm 2 (Phía dưới)
        fig.add_trace(go.Scatter(
            x=[l], y=[y_bot], mode="markers",
            marker=dict(symbol="circle", size=7, color=COLOR_SUP, line=dict(width=1, color=COLOR_SUP)),
            hoverinfo="skip"
        ))
        # Thanh thẳng đứng nối liên kết
        fig.add_trace(go.Scatter(
            x=[l, l], y=[y_top, y_bot], mode="lines",
            line=dict(color=COLOR_SUP, width=1.5),
            hoverinfo="skip"
        ))
        # Mặt sàn phẳng ngang
        fig.add_trace(go.Scatter(
            x=[l - l / 34, l + l / 34], y=[y_floor, y_floor], mode="lines",
            line=dict(color=COLOR_SUP, width=2),
            hoverinfo="skip"
        ))
    else:
        # Gối ngàm cố định (Cantilever - Giữ nguyên)
        fig.add_shape(type="rect", x0=l, x1=l + l / 42, y0=-0.42, y1=0.42,
                      fillcolor=COLOR_SUP, line={"color": COLOR_SUP})


def plot_load_diagram_single(data: BeamInput) -> go.Figure:
    l = data.length
    fig = base_figure("Load diagram", l)

    fig.update_yaxes(
        range=[-1.55, 1.25],
        fixedrange=True,
        showticklabels=False,
        title=""
    )

    # ------------------------------------------------------
    # 0) VẼ DẦM + GỐI
    # ------------------------------------------------------
    fig.add_trace(go.Scatter(
        x=[0, l], y=[0, 0],
        mode="lines",
        line={"color": COLOR_BEAM, "width": 8},
        hoverinfo="skip",
        showlegend=False
    ))
    draw_supports_single(fig, data)

    # ======================================================
    # 1) POINT LOAD
    # Quy ước:
    #   P < 0  -> tải xuống  -> mũi tên từ trên chĩa xuống dầm
    #   P > 0  -> tải lên    -> mũi tên từ dưới chĩa lên dầm
    # ======================================================
    for P, xp in data.point_loads:
        if P < 0:
            y_tip = 0.05
            y_tail = 0.78
            y_text = 0.90
        else:
            y_tip = -0.05
            y_tail = -0.78
            y_text = -0.92

        fig.add_annotation(
            x=xp, y=y_tip,
            ax=xp, ay=y_tail,
            xref="x", yref="y",
            axref="x", ayref="y",
            showarrow=True,
            arrowhead=3,
            arrowsize=1.1,
            arrowwidth=2.2,
            arrowcolor=COLOR_SFD,
            text=""
        )

        fig.add_annotation(
            x=xp,
            y=y_text,
            text=f"{P:g} kN",
            showarrow=False,
            font={"size": 11, "color": COLOR_SFD}
        )

    # ======================================================
    # 2) UDL
    # Quy ước:
    #   q < 0  -> tải xuống  -> mũi tên từ trên chĩa xuống dầm
    #   q > 0  -> tải lên    -> mũi tên từ dưới chĩa lên dầm
    # ======================================================
    for q, x1, x2 in data.udls:
        n = max(3, int(np.ceil((x2 - x1) / 0.5)) + 1)
        xs = np.linspace(x1, x2, n)

        if q < 0:
            y_load = 0.58
            y_label = 0.70

            fig.add_trace(go.Scatter(
                x=[x1, x2, x2, x1, x1],
                y=[0, 0, y_load, y_load, 0],
                fill="toself",
                mode="lines",
                line={"color": "#168f2c", "width": 1.2},
                fillcolor="rgba(22,143,44,0.16)",
                hovertemplate=f"UDL: {q:g} kN/m<extra></extra>",
                showlegend=False
            ))

            for xv in xs:
                fig.add_annotation(
                    x=xv, y=0.02,
                    ax=xv, ay=y_load,
                    xref="x", yref="y",
                    axref="x", ayref="y",
                    showarrow=True,
                    arrowhead=2,
                    arrowsize=0.95,
                    arrowwidth=1.5,
                    arrowcolor="#168f2c",
                    text=""
                )
        else:
            y_load = -0.58
            y_label = -0.72

            fig.add_trace(go.Scatter(
                x=[x1, x2, x2, x1, x1],
                y=[0, 0, y_load, y_load, 0],
                fill="toself",
                mode="lines",
                line={"color": "#168f2c", "width": 1.2},
                fillcolor="rgba(22,143,44,0.16)",
                hovertemplate=f"UDL: {q:g} kN/m<extra></extra>",
                showlegend=False
            ))

            for xv in xs:
                fig.add_annotation(
                    x=xv, y=-0.02,
                    ax=xv, ay=y_load,
                    xref="x", yref="y",
                    axref="x", ayref="y",
                    showarrow=True,
                    arrowhead=2,
                    arrowsize=0.95,
                    arrowwidth=1.5,
                    arrowcolor="#168f2c",
                    text=""
                )

        fig.add_annotation(
            x=(x1 + x2) / 2,
            y=y_label,
            text=f"{q:g} kN/m",
            showarrow=False,
            font={"size": 11, "color": "#168f2c"}
        )

    # ======================================================
    # 3) UVL
    # ======================================================
    for q, x1, x2 in data.uvls:
        n = max(6, int(np.ceil((x2 - x1) / 0.45)) + 1)
        xs = np.linspace(x1, x2, n)

        amp = 0.58 if q < 0 else -0.58

        if data.uvl_type == "increase":
            y_profile = np.linspace(0.0, amp, n)
        else:
            y_profile = np.linspace(amp, 0.0, n)

        fig.add_trace(go.Scatter(
            x=list(xs) + [xs[-1], xs[0]],
            y=list(y_profile) + [0, 0],
            fill="toself",
            mode="lines",
            line={"color": "#168f2c", "width": 1.2},
            fillcolor="rgba(22,143,44,0.16)",
            hovertemplate=f"UVL: {q:g} kN/m<extra></extra>",
            showlegend=False
        ))

        for xv, yv in zip(xs, y_profile):
            if abs(yv) < 0.06:
                continue

            if q < 0:
                fig.add_annotation(
                    x=xv, y=0.02,
                    ax=xv, ay=yv,
                    xref="x", yref="y",
                    axref="x", ayref="y",
                    showarrow=True,
                    arrowhead=2,
                    arrowsize=0.9,
                    arrowwidth=1.4,
                    arrowcolor="#168f2c",
                    text=""
                )
            else:
                fig.add_annotation(
                    x=xv, y=-0.02,
                    ax=xv, ay=yv,
                    xref="x", yref="y",
                    axref="x", ayref="y",
                    showarrow=True,
                    arrowhead=2,
                    arrowsize=0.9,
                    arrowwidth=1.4,
                    arrowcolor="#168f2c",
                    text=""
                )

        fig.add_annotation(
            x=(x1 + x2) / 2,
            y=(0.72 if q < 0 else -0.76),
            text=f"{q:g} kN/m",
            showarrow=False,
            font={"size": 11, "color": "#168f2c"}
        )

    # ======================================================
    # 4) POINT MOMENT
    # Quy ước:
    #   M < 0  -> clockwise
    #   M > 0  -> counter-clockwise
    # ======================================================
    for M, xp in data.point_moments:
        col = "#ff2f92"

        if M < 0:
            # moment âm = chiều kim đồng hồ
            fig.add_annotation(
                x=xp + 0.52, y=0.40,
                ax=xp - 0.52, ay=0.40,
                xref="x", yref="y",
                axref="x", ayref="y",
                showarrow=True,
                arrowhead=2,
                arrowsize=1.0,
                arrowwidth=2.0,
                arrowcolor=col,
                text=""
            )
            fig.add_annotation(
                x=xp - 0.52, y=0.62,
                ax=xp + 0.52, ay=0.62,
                xref="x", yref="y",
                axref="x", ayref="y",
                showarrow=True,
                arrowhead=2,
                arrowsize=1.0,
                arrowwidth=2.0,
                arrowcolor=col,
                text=""
            )
        else:
            # moment dương = ngược chiều kim đồng hồ
            fig.add_annotation(
                x=xp - 0.52, y=0.40,
                ax=xp + 0.52, ay=0.40,
                xref="x", yref="y",
                axref="x", ayref="y",
                showarrow=True,
                arrowhead=2,
                arrowsize=1.0,
                arrowwidth=2.0,
                arrowcolor=col,
                text=""
            )
            fig.add_annotation(
                x=xp + 0.52, y=0.62,
                ax=xp - 0.52, ay=0.62,
                xref="x", yref="y",
                axref="x", ayref="y",
                showarrow=True,
                arrowhead=2,
                arrowsize=1.0,
                arrowwidth=2.0,
                arrowcolor=col,
                text=""
            )

        fig.add_annotation(
            x=xp,
            y=0.86,
            text=f"{M:g} kNm",
            showarrow=False,
            font={"size": 11, "color": col}
        )

    return fig





def plot_sfd_single(result: BeamResult) -> go.Figure:
    v_range = _padded_range(result.shear)
    fig = synced_figure("Shear Force Diagram", float(result.x[-1]), y_range=v_range, y_title="Shear (kN)")
    fig.add_trace(go.Scatter(x=result.x, y=result.shear, mode="lines",
                             fill="tozeroy", line={"color": COLOR_SFD, "width": 2},
                             fillcolor="rgba(11,95,255,0.20)",
                             hovertemplate="x=%{x:.2f}m  V=%{y:.2f}kN<extra></extra>"))
    fig.update_yaxes(fixedrange=True)
    return fig


def plot_bmd_single(result: BeamResult) -> go.Figure:
    m_range = _padded_range(result.moment)
    fig = synced_figure(
        "Bending Moment Diagram",
        float(result.x[-1]),
        y_range=m_range,
        y_title="Moment (kNm)"
    )
    fig.add_trace(go.Scatter(
        x=result.x,
        y=result.moment,
        mode="lines",
        fill="tozeroy",
        line={
            "color": COLOR_BMD,
            "width": 2
        },
        fillcolor="rgba(255,43,43,0.22)",
        hovertemplate=
        "x=%{x:.2f}m  M=%{y:.2f}kNm<extra></extra>"
    ))
    fig.update_yaxes(autorange="reversed", fixedrange=True)
    return fig


def plot_elastic_single(data: BeamInput, result: BeamResult | None) -> go.Figure:
    l = data.length
    # Tạo figure mới với layout giống load diagram
    fig = base_figure("Elastic Curve", l)
    fig.update_yaxes(
        range=[-1.5, 1.2],
        fixedrange=True,
        title="Deflection (visual)"
    )
    # Vẽ dầm
    fig.add_trace(go.Scatter(
        x=[0, l], y=[0, 0],
        mode="lines",
        line={"color": COLOR_BEAM, "width": 8},
        hoverinfo="skip"
    ))
    # Vẽ gối
    draw_supports_single(fig, data)

    # Vẽ đường cong võng nếu có kết quả
    if result is not None:
        mw = float(np.max(np.abs(result.deflection)))
        y = -result.deflection * (0.65 / mw) if mw > 0 else result.deflection
        fig.add_trace(go.Scatter(
            x=result.x, y=y,
            mode="lines",
            line={"color": COLOR_ELAST, "width": 4},
            hovertemplate="x=%{x:.2f}m  w/EI=%{customdata:.4f}<extra></extra>",
            customdata=result.deflection
        ))

    fig.update_layout(title={"text": "<b>Elastic Curve</b>", "x": 0.5, "font": {"size": 15}})
    return fig


def metric_strip_single(result: BeamResult | None, data: BeamInput) -> None:
    if result is None:
        values = [("Span", f"{data.length:.2f} m"),
                  ("Point loads", str(len(data.point_loads))),
                  ("UDL / UVL", f"{len(data.udls)} / {len(data.uvls)}"),
                  ("Status", "Ready")]
    elif data.beam_type == "simple":
        idx_v = int(np.argmax(np.abs(result.shear)))
        idx_m = int(np.argmax(np.abs(result.moment)))
        idx_w = int(np.argmax(np.abs(result.deflection)))
        values = [("R1 / R2", f"{result.r1:.2f} / {result.r2:.2f} kN"),
                  ("Vmax", f"{result.shear[idx_v]:.2f} kN"),
                  ("Mmax", f"{result.moment[idx_m]:.2f} kNm"),
                  ("wmax/EI", f"{result.deflection[idx_w]:.4f}")]
    else:
        idx_v = int(np.argmax(np.abs(result.shear)))
        idx_w = int(np.argmax(np.abs(result.deflection)))
        values = [("RV", f"{result.rv_fixed:.2f} kN"),
                  ("MR", f"{result.mr_fixed:.2f} kNm"),
                  ("Vmax", f"{result.shear[idx_v]:.2f} kN"),
                  ("wmax/EI", f"{result.deflection[idx_w]:.4f}")]
    metric_html(values)

# ==========================================================
# AGGRID INPUT HELPERS
# ==========================================================
def _ensure_min_rows(df: pd.DataFrame, n: int = 8) -> pd.DataFrame:
    """Đảm bảo bảng luôn có sẵn vài dòng trống để nhập liên tục bằng bàn phím."""
    df = df.copy()
    if len(df) < n:
        extra = pd.DataFrame([{c: None for c in df.columns} for _ in range(n - len(df))])
        df = pd.concat([df, extra], ignore_index=True)
    return df


def keyboard_grid(
    key: str,
    df: pd.DataFrame,
    height: int = 260,
    min_rows: int = 8,
    reset_cnt: int = 0,
) -> pd.DataFrame:
    """
    Bảng nhập liệu bằng AgGrid.
    reset_cnt: truyền vào _get_reset_cnt(prefix) để buộc AgGrid
    tạo widget mới (xóa data browser DOM) khi New Model được bấm.
    """
    df = _ensure_min_rows(df, min_rows)

    gb = GridOptionsBuilder.from_dataframe(df)
    gb.configure_default_column(
        editable=True, resizable=True, sortable=False, filter=False,
    )
    gb.configure_grid_options(
        enterMovesDownAfterEdit=True,
        singleClickEdit=True,
        stopEditingWhenCellsLoseFocus=True,
        suppressRowClickSelection=True,
        rowSelection="single",
        enableRangeSelection=True,
        undoRedoCellEditing=True,
        undoRedoCellEditingLimit=30,
    )
    for c in df.columns:
        gb.configure_column(c, type=["numericColumn"], precision=4)

    grid_options = gb.build()

    # Nhúng reset_cnt vào key → Streamlit tạo widget mới khi counter thay đổi
    # → AgGrid render từ đầu với df rỗng → data trong browser bị clear
    widget_key = f"{key}__r{reset_cnt}"

    resp = AgGrid(
        df,
        gridOptions=grid_options,
        key=widget_key,
        height=height,
        fit_columns_on_grid_load=True,
        allow_unsafe_jscode=True,
        update_mode=GridUpdateMode.VALUE_CHANGED | GridUpdateMode.MODEL_CHANGED,
        data_return_mode=DataReturnMode.AS_INPUT,
        reload_data=False,
    )

    return pd.DataFrame(resp["data"])


def clean_rows_aggrid(df: pd.DataFrame, cols: list[str]) -> list[tuple]:
    """
    Làm sạch dữ liệu từ AgGrid:
    - bỏ các dòng trống
    - ép float
    - chỉ giữ dòng đủ dữ liệu
    """
    if df is None or df.empty:
        return []

    out = []
    for _, row in df.iterrows():
        vals = []
        ok = True
        for c in cols:
            v = row.get(c, None)
            if pd.isna(v) or v == "":
                ok = False
                break
            try:
                vals.append(float(v))
            except Exception:
                ok = False
                break
        if ok:
            out.append(tuple(vals))
    return out
def render_single_beam() -> None:
    with st.sidebar:
        st.header("⚙️ Single Beam — Input")
        if st.button("🆕 New Model", type="primary", use_container_width=True, key="sb_new_btn"):
            # Quét sạch MỌI biến, widget, data, result có tiền tố "sb_"
            reset_keys_with_prefix("sb_")
            st.rerun()
            for k in [
                "sb_pl_df", "sb_pm_df", "sb_udl_df", "sb_uvl_df",
                "sb_pl_grid", "sb_pm_grid", "sb_udl_grid", "sb_uvl_grid",
                "sb_result", "sb_input"
            ]:
                st.session_state.pop(k, None)
            st.rerun()
        st.divider()
        length = st.number_input("Chiều dài L (m)", min_value=0.01, value=10.0, step=0.5,
                                 format="%.2f", key="sb_L")
        bt = st.radio("Type of Beam", ["Simply Supported", "Cantilever"],
                      horizontal=True, key="sb_bt")
        ut = st.radio("UVL Type", ["Increase", "Decrease"],
                      horizontal=True, key="sb_ut")
        st.divider()
        st.caption("Nhập tải trong các bảng ở vùng làm việc chính.")

    data = BeamInput(
        length=float(length),
        beam_type="simple" if bt == "Simply Supported" else "cantilever",
        uvl_type="increase" if ut == "Increase" else "decrease",
    )

    # ── Load tables: keyboard-friendly via AgGrid ─────────────────────
    pl_default = pd.DataFrame(columns=["P (kN)", "x (m)"])
    pm_default = pd.DataFrame(columns=["M (kNm)", "x (m)"])
    udl_default = pd.DataFrame(columns=["q (kN/m)", "x1 (m)", "x2 (m)"])
    uvl_default = pd.DataFrame(columns=["qmax (kN/m)", "x1 (m)", "x2 (m)"])

    # Chỉ khởi tạo nếu chưa có (setdefault an toàn với DataFrame)
    if "sb_pl_df"  not in st.session_state: st.session_state["sb_pl_df"]  = pl_default.copy()
    if "sb_pm_df"  not in st.session_state: st.session_state["sb_pm_df"]  = pm_default.copy()
    if "sb_udl_df" not in st.session_state: st.session_state["sb_udl_df"] = udl_default.copy()
    if "sb_uvl_df" not in st.session_state: st.session_state["sb_uvl_df"] = uvl_default.copy()

    _sb_cnt = _get_reset_cnt("sb_")

    t1, t2, t3, t4 = st.tabs(["Point Load", "Point Moment", "UDL", "UVL"])

    with t1:
        st.caption("Mẹo: click 1 ô rồi dùng Tab để nhập nhanh.")
        pl = keyboard_grid("sb_pl_grid", st.session_state["sb_pl_df"], height=220, min_rows=8, reset_cnt=_sb_cnt)
        st.session_state["sb_pl_df"] = pl

    with t2:
        st.caption("Mẹo: click 1 ô rồi dùng Tab để nhập nhanh.")
        pm = keyboard_grid("sb_pm_grid", st.session_state["sb_pm_df"], height=220, min_rows=8, reset_cnt=_sb_cnt)
        st.session_state["sb_pm_df"] = pm

    with t3:
        st.caption("Mẹo: click 1 ô rồi dùng Tab để nhập nhanh.")
        udl = keyboard_grid("sb_udl_grid", st.session_state["sb_udl_df"], height=220, min_rows=8, reset_cnt=_sb_cnt)
        st.session_state["sb_udl_df"] = udl

    with t4:
        st.caption("Mẹo: click 1 ô rồi dùng ↑ ↓ ← →, Enter, Tab để nhập nhanh.")
        uvl = keyboard_grid("sb_uvl_grid", st.session_state["sb_uvl_df"], height=220, min_rows=8, reset_cnt=_sb_cnt)
        st.session_state["sb_uvl_df"] = uvl

    data.point_loads = clean_rows_aggrid(pl, ["P (kN)", "x (m)"])
    data.point_moments = clean_rows_aggrid(pm, ["M (kNm)", "x (m)"])
    data.udls = clean_rows_aggrid(udl, ["q (kN/m)", "x1 (m)", "x2 (m)"])
    data.uvls = clean_rows_aggrid(uvl, ["qmax (kN/m)", "x1 (m)", "x2 (m)"])

    result: BeamResult | None = st.session_state.get("sb_result")

    if st.button("▶ Solve", type="primary", use_container_width=True, key="sb_solve"):
        errs = validate_single(data)
        if errs:
            for e in errs: st.error(e)
            result = None
        else:
            try:
                result = solve_beam(data)
                st.session_state.sb_result = result
                st.session_state.sb_input  = data
            except Exception as e:
                st.error(str(e)); result = None

    metric_strip_single(result, data)

    left, right = st.columns([1.7, 1], gap="large")
    with left:
        a, b = st.columns(2)
        with a: st.plotly_chart(plot_load_diagram_single(data), use_container_width=True)
        with b: st.plotly_chart(
            plot_sfd_single(result) if result else base_figure("Shear Force Diagram", data.length, "kN"),
            use_container_width=True)
        c, d = st.columns(2)
        with c: st.plotly_chart(
            plot_bmd_single(result) if result else base_figure("Moment Diagram", data.length, "kNm"),
            use_container_width=True)
        with d: st.plotly_chart(plot_elastic_single(data, result), use_container_width=True)


    with right:
        figures = []

        if result:
            figures.append(
            ("Load Diagram",
             plot_load_diagram_single(data))
        )

            figures.append(
            ("Shear Force Diagram",
             plot_sfd_single(result))
        )

            figures.append(
            ("Bending Moment Diagram",
             plot_bmd_single(result))
        )

            figures.append(
            ("Elastic Curve",
             plot_elastic_single(data, result))
        )

    report_panel(
        result.report if result else None,
        "Thuyết Minh — Dầm Đơn",
        "single_beam",
        figures=figures
    )
# ══════════════════════════════════════════════════════
#  ── TAB 2: CONTINUOUS BEAM ──────────────────────────
# ══════════════════════════════════════════════════════

def render_continuous_beam() -> None:
    with st.sidebar:
        st.header("⚙️ Continuous Beam — Input")
        if st.button("🆕 New Model", type="primary", use_container_width=True, key="cb_new_btn"):
            reset_keys_with_prefix("cb_")
            st.rerun()
        st.divider()

        n_spans = st.number_input("Số nhịp", min_value=1, max_value=20, value=2, step=1, key="cb_nspans")
        st.divider()

        st.markdown("**Thông số từng nhịp**")
        span_lengths, span_EIs = [], []
        for i in range(int(n_spans)):
            c1, c2 = st.columns(2)
            with c1:
                L_i = st.number_input(f"L{i+1} (m)", min_value=0.01, value=5.0, step=0.5, format="%.2f", key=f"cb_L{i}")
            with c2:
                EI_i = st.number_input(f"EI{i+1}", min_value=1e-6, value=1.0, step=100.0, format="%.4g", key=f"cb_EI{i}")
            span_lengths.append(float(L_i))
            span_EIs.append(float(EI_i))

        st.divider()
        n_nodes_boundary = int(n_spans) + 1
        st.markdown("**Gối đỡ**")
        support_kinds = []
        for i in range(n_nodes_boundary):
            xpos = sum(span_lengths[:i])
            kind = st.selectbox(
                f"Node {i} (x={xpos:.2f}m)",
                ["pin", "roller", "fixed", "free"],
                key=f"cb_sup{i}",
                index=0 if i == 0 else (0 if i == n_nodes_boundary-1 else 0),
            )
            support_kinds.append(kind)

        st.divider()
        st.caption("Nhập tải trọng trong bảng ở vùng làm việc chính.")

    st.markdown("#### Tải trọng từng nhịp")
    span_pl, span_udl, span_pm = [], [], []
    _cb_cnt = _get_reset_cnt("cb_")

    for i in range(int(n_spans)):
        with st.expander(f"Nhịp {i+1}  (L = {span_lengths[i]:.2f} m)", expanded=(i == 0)):
            t1, t2, t3 = st.tabs(["Point Load", "UDL", "Point Moment"])
            with t1:
                key_pl = f"cb_pl_df_{i}"
                if key_pl not in st.session_state:
                    st.session_state[key_pl] = pd.DataFrame(columns=["P (kN)", "x_local (m)"])
                df_pl = keyboard_grid(f"cb_pl_grid_{i}", st.session_state[key_pl], height=210, min_rows=6, reset_cnt=_cb_cnt)
                st.session_state[key_pl] = df_pl

            with t2:
                key_udl = f"cb_udl_df_{i}"
                if key_udl not in st.session_state:
                    st.session_state[key_udl] = pd.DataFrame(columns=["q (kN/m)", "x1_local (m)", "x2_local (m)"])
                df_udl = keyboard_grid(f"cb_udl_grid_{i}", st.session_state[key_udl], height=210, min_rows=6, reset_cnt=_cb_cnt)
                st.session_state[key_udl] = df_udl

            with t3:
                key_pm = f"cb_pm_df_{i}"
                if key_pm not in st.session_state:
                    st.session_state[key_pm] = pd.DataFrame(columns=["M (kNm)", "x_local (m)"])
                df_pm = keyboard_grid(f"cb_pm_grid_{i}", st.session_state[key_pm], height=210, min_rows=6, reset_cnt=_cb_cnt)
                st.session_state[key_pm] = df_pm

            span_pl.append(clean_rows_aggrid(df_pl, ["P (kN)", "x_local (m)"]))
            span_udl.append(clean_rows_aggrid(df_udl, ["q (kN/m)", "x1_local (m)", "x2_local (m)"]))
            span_pm.append(clean_rows_aggrid(df_pm, ["M (kNm)", "x_local (m)"]))

    errors = []

    errors.extend(
        validate_point_loads(
            span_pl[i],
            span_lengths[i]
        )
    )

    errors.extend(
        validate_udls(
            span_udl[i],
            span_lengths[i]
        )
    )

    errors.extend(
        validate_point_moments(
            span_pm[i],
            span_lengths[i]
        )
    )

    if errors:

        st.error(
            f"❌ Nhịp {i + 1}: Có tải trọng nằm ngoài chiều dài nhịp."
        )

        st.info(
            f"📏 Chỉ được nhập trong khoảng: "
            f"0.00 ≤ x_local ≤ {span_lengths[i]:.2f} m"
        )

        for msg in errors:
            st.warning(msg)
    result_cb: ContinuousBeamResult | None = st.session_state.get("cb_result")
    disable_solve = False

    for i in range(int(n_spans)):
        errors = []

        errors.extend(
            validate_point_loads(
                span_pl[i],
                span_lengths[i]
            )
        )

        errors.extend(
            validate_udls(
                span_udl[i],
                span_lengths[i]
            )
        )

        errors.extend(
            validate_point_moments(
                span_pm[i],
                span_lengths[i]
            )
        )

        if errors:

            st.error(
                f"❌ Nhịp {i + 1}: Có tải trọng không hợp lệ."
            )

            st.info(
                f"📏 Chỉ được nhập trong khoảng "
                f"0 ≤ x_local ≤ {span_lengths[i]:.2f} m"
            )

            for err in errors:
                st.warning(err)
        if validate_point_loads(span_pl[i], span_lengths[i]):
            disable_solve = True

        if validate_udls(span_udl[i], span_lengths[i]):
            disable_solve = True

        if validate_point_moments(span_pm[i], span_lengths[i]):
            disable_solve = True
    solve_clicked = st.button(
        "▶ Solve",
        type="primary",
        use_container_width=True,
        key="cb_solve",
        disabled=disable_solve,
    )

    if solve_clicked:
        try:
            spans_def = []

            for i in range(int(n_spans)):
                spans_def.append(
                    SpanDef(
                        length=span_lengths[i],
                        EI=span_EIs[i],
                        point_loads=[(P, x) for P, x in span_pl[i]],
                        udls=[(q, x1, x2) for q, x1, x2 in span_udl[i]],
                        point_moments=[(M, x) for M, x in span_pm[i]],
                    )
                )

            supports_def = [
                SupportDef(node=i, kind=support_kinds[i])
                for i in range(n_nodes_boundary)
                if support_kinds[i] != "free"
            ]

            cb_input = ContinuousBeamInput(
                spans=spans_def,
                supports=supports_def,
            )

            result_cb = solve_continuous_beam(cb_input)

            st.session_state.cb_result = result_cb
            st.session_state.cb_input = cb_input

        except Exception as e:

            st.error(f"Lỗi tính toán: {e}")

            result_cb = None

    if result_cb is None:
        total_L = sum(span_lengths)
        metric_html([("Tổng L", f"{total_L:.2f} m"), ("Số nhịp", str(n_spans)), ("Số gối", str(sum(1 for k in support_kinds if k != "free"))), ("Status", "Ready")])
    else:
        xv = result_cb.x_global
        V, M, w = result_cb.shear, result_cb.moment, result_cb.deflection
        iv, im, iw = int(np.argmax(np.abs(V))), int(np.argmax(np.abs(M))), int(np.argmax(np.abs(w)))
        metric_html([("Vmax", f"{V[iv]:.3f} kN  @x={xv[iv]:.2f}m"),
                     ("Mmax", f"{M[im]:.3f} kNm @x={xv[im]:.2f}m"),
                     ("wmax/EI", f"{w[iw]:.5f} m  @x={xv[iw]:.2f}m"),
                     ("Gối", f"{len(result_cb.reactions)} phản lực")])

    total_L_plot = sum(span_lengths)
    fig_load = None
    fig_sfd = None
    fig_bmd = None
    fig_el = None
    left, right = st.columns([1.7, 1], gap="large")
    with left:
        fig_load = _cb_load_diagram(span_lengths, span_EIs, span_pl, span_udl, support_kinds,span_pm,
    )
        a, b = st.columns(2)
        with a: st.plotly_chart(fig_load, use_container_width=True)
        with b:
            if result_cb:
                v_range = _padded_range(result_cb.shear)
                fig_sfd = synced_figure("Shear Force Diagram", total_L_plot, y_range=v_range, y_title="V (kN)")
                fig_sfd.add_trace(go.Scatter(x=result_cb.x_global, y=result_cb.shear, mode="lines", fill="tozeroy", line={"color": COLOR_SFD, "width": 2}, fillcolor="rgba(11,95,255,0.20)", hovertemplate="x=%{x:.3f}m  V=%{y:.3f}kN<extra></extra>"))
                fig_sfd.update_yaxes(fixedrange=True)
                st.plotly_chart(fig_sfd, use_container_width=True)
            else:
                st.plotly_chart(base_figure("Shear Force Diagram", total_L_plot, "V (kN)"), use_container_width=True)

        c, d = st.columns(2)
        with c:
            if result_cb:
                m_range = _padded_range(result_cb.moment)
                fig_bmd = synced_figure("Bending Moment Diagram", total_L_plot, y_range=m_range, y_title="M (kNm)")
                fig_bmd.add_trace(go.Scatter(x=result_cb.x_global, y=result_cb.moment, mode="lines", fill="tozeroy", line={"color": COLOR_BMD, "width": 2}, fillcolor="rgba(255,43,43,0.22)", hovertemplate="x=%{x:.3f}m  M=%{y:.3f}kNm<extra></extra>"))
                fig_bmd.update_yaxes(autorange="reversed", fixedrange=True)
                st.plotly_chart(fig_bmd, use_container_width=True)
            else:
                st.plotly_chart(base_figure("Bending Moment Diagram", total_L_plot, "M (kNm)"), use_container_width=True)
        with d:
            if result_cb:
                # 1. Khởi dựng phôi nền sạch từ hàm bổ trợ (Đồng bộ dải hiển thị cố định)
                fig_el = _cb_draw_base_beam_and_supports(total_L_plot, span_lengths, support_kinds)
                fig_el.update_layout(title={"text": "<b>Elastic Curve</b>"})

                # 2. Vẽ đường cong võng (khống chế biên độ trực quan tối đa là 0.65 để thoáng đồ thị)
                mw = float(np.max(np.abs(result_cb.deflection))) + 1e-30
                y_vis = -result_cb.deflection * (0.65 / mw)

                fig_el.add_trace(go.Scatter(
                    x=result_cb.x_global, y=y_vis, mode="lines",
                    line={"color": COLOR_ELAST, "width": 3},
                    hovertemplate="x=%{x:.3f}m  w/EI=%{customdata:.5f}<extra></extra>",
                    customdata=result_cb.deflection
                ))

                # 3. Đồng bộ trục Y khớp hoàn toàn với biểu đồ Load Diagram
                fig_el.update_yaxes(
                    range=[-1.5, 1.2],
                    fixedrange=True,
                    showticklabels=True,
                    title="Deflection (visual)"
                )
                st.plotly_chart(fig_el, use_container_width=True)
            else:
                fig_empty = go.Figure()
                fig_empty.update_layout(height=315, title="<b>Elastic Curve</b>", xaxis=dict(range=[0, total_L_plot]))
                fig_empty.update_yaxes(range=[-1.5, 1.2], fixedrange=True)
                st.plotly_chart(fig_empty, use_container_width=True)
    with right:

        figures = []

        if result_cb:

            if fig_load is not None:
                figures.append(("Load Diagram", fig_load))

            if fig_sfd is not None:
                figures.append(("Shear Force Diagram", fig_sfd))

            if fig_bmd is not None:
                figures.append(("Bending Moment Diagram", fig_bmd))

            if fig_el is not None:
                figures.append(("Elastic Curve", fig_el))

        report_panel(
            result_cb.report if result_cb else None,
            "Thuyết Minh — Dầm Liên Tục",
            "cont_beam",
            figures=figures
        )

def _cb_draw_base_beam_and_supports(total_L, span_lengths, support_kinds) -> go.Figure:
    """Hàm dựng khung dầm và gối đỡ nền đồng bộ tỷ lệ hiển thị"""

    fig = go.Figure()

    # khoảng trống hai đầu để không cắt gối
    margin_x = max(total_L * 0.06, 0.5)

    # 1. Layout
    fig.update_layout(
        title=dict(
            x=0.5,
            xanchor="center",
            font=dict(size=14, color="#333333")
        ),

        height=315,

        margin=dict(
            l=55,
            r=20,
            t=60,
            b=75
        ),

        paper_bgcolor="rgba(0,0,0,0)",
        plot_bgcolor="rgba(0,0,0,0)",
        showlegend=False,

        xaxis=dict(
            title="x (m)",
            range=[-margin_x, total_L + margin_x],
            showgrid=True,
            linecolor="gray",
            gridcolor="rgba(128,128,128,0.15)"
        )
    )


    fig.update_yaxes(
        range=[-1.5, 1.2],
        fixedrange=True,
        showgrid=False,
        linecolor="gray",
        showticklabels=False,
        title=""
    )


    # 2. Thanh dầm
    fig.add_trace(
        go.Scatter(
            x=[0, total_L],
            y=[0, 0],
            mode="lines",
            line={
                "color": COLOR_BEAM,
                "width": 8
            },
            hoverinfo="skip"
        )
    )


    # Nhãn nhịp
    x_acc = 0.0

    for i, Ls in enumerate(span_lengths):
        mid = x_acc + Ls / 2

        fig.add_annotation(
            x=mid,
            y=0.2,
            text=f"L{i+1}={Ls:.1f}m",
            showarrow=False,
            font={
                "size":10,
                "color":"#555"
            }
        )

        x_acc += Ls



    # 3. Gối
    node_xs = [0.0] + list(np.cumsum(span_lengths))

    support_size = max(total_L * 0.035, 0.25)


    for i, kind in enumerate(support_kinds):

        xp = node_xs[i]


        # gối pin
        if kind == "pin":

            fig.add_trace(
                go.Scatter(
                    x=[
                        xp,
                        xp + support_size,
                        xp - support_size,
                        xp
                    ],

                    y=[
                        0,
                        -0.37,
                        -0.37,
                        0
                    ],

                    fill="toself",
                    mode="lines",
                    line={
                        "color":COLOR_SUP,
                        "width":1.5
                    },
                    fillcolor=COLOR_SUP,
                    hoverinfo="skip"
                )
            )


        # gối con lăn
        elif kind == "roller":

            y_top = -0.08
            y_bot = -0.36
            y_floor = -0.44


            fig.add_trace(
                go.Scatter(
                    x=[xp],
                    y=[y_top],
                    mode="markers",
                    marker=dict(
                        symbol="circle",
                        size=7,
                        color=COLOR_SUP
                    ),
                    hoverinfo="skip"
                )
            )


            fig.add_trace(
                go.Scatter(
                    x=[xp],
                    y=[y_bot],
                    mode="markers",
                    marker=dict(
                        symbol="circle",
                        size=7,
                        color=COLOR_SUP
                    ),
                    hoverinfo="skip"
                )
            )


            fig.add_trace(
                go.Scatter(
                    x=[xp,xp],
                    y=[y_top,y_bot],
                    mode="lines",
                    line=dict(
                        color=COLOR_SUP,
                        width=1.5
                    ),
                    hoverinfo="skip"
                )
            )


            fig.add_trace(
                go.Scatter(
                    x=[
                        xp-support_size,
                        xp+support_size
                    ],
                    y=[
                        y_floor,
                        y_floor
                    ],
                    mode="lines",
                    line=dict(
                        color=COLOR_SUP,
                        width=2
                    ),
                    hoverinfo="skip"
                )
            )



        # ngàm
        elif kind == "fixed":

            fig.add_shape(
                type="rect",
                x0=xp-total_L/80,
                x1=xp+total_L/80,
                y0=-0.42,
                y1=0.42,
                fillcolor=COLOR_SUP,
                line={"color":COLOR_SUP}
            )


        fig.add_annotation(
            x=xp,
            y=-0.68,
            text=f"N{i}",
            showarrow=False,
            font={
                "size":10,
                "color":COLOR_SUP
            }
        )


    return fig


def _cb_load_diagram(span_lengths, span_EIs, span_pl, span_udl, support_kinds, span_pm=None) -> go.Figure:

    total_L = sum(span_lengths)

    fig = _cb_draw_base_beam_and_supports(
        total_L,
        span_lengths,
        support_kinds
    )

    fig.update_layout(
        title={
            "text": "<b>Load Diagram — Dầm liên tục</b>",
            "x":0.5
        }
    )

    node_xs = [0.0] + list(np.cumsum(span_lengths))


    # ===============================
    # TẢI TRỌNG THEO TỪNG NHỊP
    # ===============================
    for i in range(len(span_lengths)):

        x0_span = node_xs[i]


        # -------------------------------
        # UDL
        # -------------------------------
        for q, x1_local, x2_local in span_udl[i]:

            if abs(q) < 1e-9:
                continue

            x1 = x0_span + x1_local
            x2 = x0_span + x2_local


            # q dương: hướng xuống vào dầm
            # q âm: hướng lên vào dầm
            if q > 0:
                y_load = -0.58
                arrow_start = -0.75
                arrow_end = -0.06
            else:
                y_load = 0.58
                arrow_start = 0.75
                arrow_end = 0.06
            # vùng tải
            fig.add_trace(
                go.Scatter(
                    x=[x1,x2,x2,x1,x1],
                    y=[0,0,y_load,y_load,0],
                    fill="toself",
                    mode="lines",
                    line=dict(
                        color="#28a745",
                        width=1
                    ),
                    fillcolor="rgba(40,167,69,0.15)",
                    hoverinfo="skip"
                )
            )


            # mũi tên UDL
            for xx in np.linspace(
                x1,
                x2,
                max(3,int((x2-x1)/0.5))
            ):

                fig.add_annotation(
                    x=xx,
                    y=arrow_end,
                    ax=xx,
                    ay=arrow_start,
                    xref="x",
                    yref="y",
                    axref="x",
                    ayref="y",
                    showarrow=True,
                    arrowhead=2,
                    arrowsize=1,
                    arrowwidth=1.5,
                    arrowcolor="#28a745"
                )


            fig.add_annotation(
                x=(x1+x2)/2,
                y=y_load*1.25,
                text=f"{q:.1f} kN/m",
                showarrow=False,
                font=dict(
                    size=10,
                    color="#28a745"
                )
            )



        # -------------------------------
        # POINT LOAD
        # -------------------------------
        for P,x_local in span_pl[i]:

            xp = x0_span + x_local


            if P > 0:
                ay = -0.75
                y = -0.05
            else:
                ay = 0.75
                y = 0.05


            fig.add_annotation(
                x=xp,
                y=y,
                ax=xp,
                ay=ay,
                xref="x",
                yref="y",
                axref="x",
                ayref="y",
                showarrow=True,
                arrowhead=3,
                arrowsize=1.1,
                arrowwidth=2,
                arrowcolor="#0b5fff"
            )

            fig.add_annotation(
                x=xp,
                y=ay,
                text=f"{P:.1f} kN",
                showarrow=False,
                font=dict(
                    size=10,
                    color="#0b5fff"
                )
            )

        # -------------------------------
        # POINT MOMENT (ký hiệu moment chuẩn)
        # -------------------------------
        if span_pm is not None:

            for M, x_local in span_pm[i]:

                xm = x0_span + x_local

                r = total_L * 0.025  # bán kính vòng moment
                y0 = 0.25

                # góc tạo vòng cung — ĐỒNG BỘ với bản MATLAB gốc:
                # M > 0 → quét CCW (dưới-phải → trên-trái, mũi tên chỉ sang trái)
                # M < 0 → quét CW  (dưới-trái → trên-phải, mũi tên chỉ sang phải)
                if M > 0:
                    theta = np.linspace(-np.pi * 0.25, np.pi * 0.85, 40)
                else:
                    theta = np.linspace(np.pi * 1.25, np.pi * 0.15, 40)

                x_arc = xm + r * np.cos(theta)
                y_arc = y0 + r * np.sin(theta)

                # vòng cung
                fig.add_trace(
                    go.Scatter(
                        x=x_arc,
                        y=y_arc,
                        mode="lines",
                        line=dict(
                            color="#ff2b8a",
                            width=2.5
                        ),
                        hoverinfo="skip"
                    )
                )

                # MŨI TÊN MOMENT ĐỒNG BỘ VỚI TẢI LỰC
                # ===============================

                xe = x_arc[-1]
                ye = y_arc[-1]

                # vector tiếp tuyến
                tx = x_arc[-1] - x_arc[-3]
                ty = y_arc[-1] - y_arc[-3]

                # chuẩn hóa vector
                length = math.sqrt(tx ** 2 + ty ** 2)

                tx /= length
                ty /= length

                # vector pháp tuyến để tạo tam giác
                nx = -ty
                ny = tx

                # SỬA LẠI KÍCH THƯỚC MŨI TÊN DẦM LIÊN TỤC TRÙNG PHIÊN BẢN CŨ
                arrow_len = total_L * 0.018
                arrow_w = total_L * 0.005

                p1 = (xe, ye)

                p2 = (
                    xe - tx * arrow_len + nx * arrow_w,
                    ye - ty * arrow_len + ny * arrow_w
                )

                p3 = (
                    xe - tx * arrow_len - nx * arrow_w,
                    ye - ty * arrow_len - ny * arrow_w
                )

                fig.add_trace(
                    go.Scatter(
                        x=[p1[0], p2[0], p3[0], p1[0]],
                        y=[p1[1], p2[1], p3[1], p1[1]],
                        fill="toself",
                        mode="lines",
                        line=dict(
                            color="#ff2b8a",
                            width=1
                        ),
                        fillcolor="#ff2b8a",
                        hoverinfo="skip"
                    )
                )
                # trị số moment
                fig.add_annotation(
                    x=xm,
                    y=y0 + r * 1.8,
                    text=f"{M:.1f} kNm",
                    showarrow=False,
                    font=dict(
                        size=10,
                        color="#ff2b8a"
                    )
                )
    return fig
# ══════════════════════════════════════════════════════
#  ── TAB 3: PLANE FRAME ──────────────────────────────
# ══════════════════════════════════════════════════════

def render_plane_frame() -> None:
    # ==========================================
    # GIAO DIỆN COMING SOON (ĐANG PHÁT TRIỂN)
    # ==========================================
    st.title("🏗️ Plane Frame (Khung Phẳng)")
    st.success(
        "**Plane Frame đã sẵn sàng tính toán.** Kết quả FEM được đối chiếu bằng nhân biểu đồ / Vereshchagin trong thuyết minh."
    )

    # Lệnh return này sẽ chặn không cho chạy đoạn code giao diện cũ ở bên dưới,
    # giúp bạn giữ nguyên được source code mà không cần phải xóa hay comment (#) từng dòng.
    st.caption("FEM 2D khung phẳng, kèm đối chiếu chuyển vị bằng nhân biểu đồ / Vereshchagin.")
    # ==========================================

    # --- TOÀN BỘ CODE CŨ DƯỚI ĐÂY ĐƯỢC GIỮ NGUYÊN ---
    with st.sidebar:
        st.header("⚙️ Plane Frame — Input")
        if st.button("🆕 New Model", type="primary", use_container_width=True, key="pf_new_btn"):
            # Quét sạch toàn bộ dữ liệu, bảng, kết quả của Tab Plane Frame
            reset_keys_with_prefix("pf_")
            st.rerun()
            for k in list(st.session_state.keys()):
                if k.startswith("pf_"): st.session_state.pop(k, None)
            st.rerun()
        st.divider()

    cfg = {"width": "stretch", "num_rows": "dynamic", "hide_index": True}
    tab_nd, tab_el, tab_sup, tab_pl_nd, _ = st.tabs(
        ["🔵 Nodes", "📐 Elements", "🔒 Supports", "⬇️ Node Loads", "📏 Element UDL"])

    with tab_nd:
        nodes_default = pd.DataFrame({"x (m)": [0.0, 0.0, 5.0, 5.0], "y (m)": [0.0, 4.0, 4.0, 0.0]})
        df_nodes = safe_data_editor("pf_nd_ed", nodes_default, **cfg)
    with tab_el:
        elems_default = pd.DataFrame(
            {"i": [0, 1, 3], "j": [1, 2, 2], "E": [200e6] * 3, "A": [0.01] * 3, "I": [1e-4] * 3,
             "udl_local": [0.0] * 3})
        df_el = safe_data_editor("pf_el_ed", elems_default, **cfg)
    with tab_sup:
        sups_default = pd.DataFrame({"node": [0, 3], "ux": [True, True], "uy": [True, True], "rz": [True, True]})
        df_sup = safe_data_editor("pf_sup_ed", sups_default, **cfg)
    with tab_pl_nd:
        nloads_default = pd.DataFrame(
            {"node": pd.Series(dtype=int), "Fx (kN)": pd.Series(dtype=float), "Fy (kN)": pd.Series(dtype=float),
             "Mz (kNm)": pd.Series(dtype=float)})
        df_nload = safe_data_editor("pf_nl_ed", nloads_default, **cfg)

    result_pf: PlaneFrameResult | None = st.session_state.get("pf_result")

    if st.button("▶ Solve", type="primary", use_container_width=True, key="pf_solve"):
        try:
            nodes = [FrameNode(x=float(r["x (m)"]), y=float(r["y (m)"])) for _, r in df_nodes.iterrows() if
                     pd.notna(r.get("x (m)")) and pd.notna(r.get("y (m)"))]
            elems = [
                FrameElement(i_node=int(r["i"]), j_node=int(r["j"]), E=float(r["E"]), A=float(r["A"]), I=float(r["I"]),
                             udl_local=float(r.get("udl_local", 0) or 0)) for _, r in df_el.iterrows() if
                pd.notna(r.get("i"))]
            sups = [
                FrameSupport(node=int(r["node"]), ux_fixed=bool(r.get("ux", True)), uy_fixed=bool(r.get("uy", True)),
                             rz_fixed=bool(r.get("rz", True))) for _, r in df_sup.iterrows() if pd.notna(r.get("node"))]
            pls = [FramePointLoad(node=int(r["node"]), Fx=float(r.get("Fx (kN)", 0) or 0),
                                  Fy=float(r.get("Fy (kN)", 0) or 0), Mz=float(r.get("Mz (kNm)", 0) or 0)) for _, r in
                   df_nload.iterrows() if pd.notna(r.get("node"))]

            pf_input = PlaneFrameInput(nodes=nodes, elements=elems, supports=sups, point_loads=pls)
            st.session_state.pf_result = solve_plane_frame(pf_input)
            st.rerun()
        except Exception as e:
            st.error(f"Lỗi: {e}")

    if result_pf is None:
        metric_html([("Nodes", str(len(df_nodes))), ("Elements", str(len(df_el))), ("Supports", str(len(df_sup))),
                     ("Status", "Ready")])
    else:
        all_V = np.concatenate([er.shear for er in result_pf.element_results])
        all_M = np.concatenate([er.moment for er in result_pf.element_results])
        all_N = np.concatenate([er.axial for er in result_pf.element_results])
        if getattr(result_pf, "vereshchagin_checks", None):
            cmp_df = pd.DataFrame([
                {
                    "Node": c["node"],
                    "DOF": c["component"],
                    "FEM": c["fem"],
                    "Vereshchagin": c["vereshchagin"],
                    "Diff": c["diff"],
                }
                for c in result_pf.vereshchagin_checks
            ])
            st.subheader("Đối chiếu FEM - Vereshchagin")
            st.dataframe(cmp_df, use_container_width=True, hide_index=True)
        metric_html([
            ("Nmax/Nmin", f"{np.max(all_N):.2f}/{np.min(all_N):.2f} kN"),
            ("Vmax", f"{np.max(np.abs(all_V)):.3f} kN"),
            ("Mmax", f"{np.max(np.abs(all_M)):.3f} kNm"),
            ("Reactions", f"{len(result_pf.reactions)} gối"),
        ])

    left, right = st.columns([1.7, 1], gap="large")
    with left:
        a, b = st.columns(2)
        with a: st.plotly_chart(_pf_geometry_plot(df_nodes, df_el, df_sup, result_pf), use_container_width=True)
        with b: st.plotly_chart(_pf_diagram_plot(result_pf, "moment", "BMD"), use_container_width=True)
        c, d = st.columns(2)
        with c: st.plotly_chart(_pf_diagram_plot(result_pf, "shear", "SFD"), use_container_width=True)
        with d: st.plotly_chart(_pf_diagram_plot(result_pf, "axial", "AFD"), use_container_width=True)
    with right:

        figures = []

        if result_pf:
            fig_geom = _pf_geometry_plot(
                df_nodes,
                df_el,
                df_sup,
                result_pf
            )

            fig_bmd = _pf_diagram_plot(
                result_pf,
                "moment",
                "BMD"
            )

            fig_sfd = _pf_diagram_plot(
                result_pf,
                "shear",
                "SFD"
            )

            fig_afd = _pf_diagram_plot(
                result_pf,
                "axial",
                "AFD"
            )

            figures.extend([
                ("Geometry", fig_geom),
                ("Bending Moment Diagram", fig_bmd),
                ("Shear Force Diagram", fig_sfd),
                ("Axial Force Diagram", fig_afd),
            ])

        report_panel(
            result_pf.report if result_pf else None,
            "Thuyết Minh — Khung Phẳng",
            "plane_frame",
            figures=figures
        )
def _pf_geometry_plot(
    df_nodes: pd.DataFrame,
    df_el: pd.DataFrame,
    df_sup: pd.DataFrame,
    result_pf=None,
) -> go.Figure:
    """
    Vẽ hình học khung phẳng 2D.

    Parameters
    ----------
    df_nodes : DataFrame
        Bảng Nodes gồm:
        - x (m)
        - y (m)

    df_el : DataFrame
        Bảng Elements gồm:
        - i
        - j

    df_sup : DataFrame
        Bảng Supports gồm:
        - node
        - ux
        - uy
        - rz

    result_pf : PlaneFrameResult, optional
        Kết quả FEM. Hiện tại dùng để giữ tương thích
        với giao diện và có thể mở rộng sau này.
    """

    fig = go.Figure()

    # ==========================================================
    # 1. KIỂM TRA DỮ LIỆU NODES
    # ==========================================================

    if df_nodes is None or df_nodes.empty:
        fig.update_layout(
            title="Plane Frame Geometry",
            xaxis_title="X (m)",
            yaxis_title="Y (m)",
            template="plotly_white",
        )
        return fig

    nodes = df_nodes.copy()

    # Đảm bảo tên cột tồn tại
    required_node_cols = ["x (m)", "y (m)"]

    if not all(c in nodes.columns for c in required_node_cols):
        fig.update_layout(
            title="Plane Frame Geometry — Invalid Node Data",
            template="plotly_white",
        )
        return fig

    # Ép kiểu số
    nodes["x (m)"] = pd.to_numeric(
        nodes["x (m)"],
        errors="coerce"
    )

    nodes["y (m)"] = pd.to_numeric(
        nodes["y (m)"],
        errors="coerce"
    )

    nodes = nodes.dropna(
        subset=["x (m)", "y (m)"]
    ).reset_index(drop=True)

    # ==========================================================
    # 2. VẼ ELEMENTS
    # ==========================================================

    if df_el is not None and not df_el.empty:

        elements = df_el.copy()

        if "i" in elements.columns and "j" in elements.columns:

            for _, row in elements.iterrows():

                try:
                    i = int(row["i"])
                    j = int(row["j"])

                    if (
                        i < 0
                        or j < 0
                        or i >= len(nodes)
                        or j >= len(nodes)
                    ):
                        continue

                    xi = float(nodes.loc[i, "x (m)"])
                    yi = float(nodes.loc[i, "y (m)"])

                    xj = float(nodes.loc[j, "x (m)"])
                    yj = float(nodes.loc[j, "y (m)"])

                    fig.add_trace(
                        go.Scatter(
                            x=[xi, xj],
                            y=[yi, yj],
                            mode="lines",
                            line=dict(
                                color="#1f77b4",
                                width=4,
                            ),
                            hovertemplate=(
                                f"Element: {i} → {j}"
                                "<br>X: %{x:.3f} m"
                                "<br>Y: %{y:.3f} m"
                                "<extra></extra>"
                            ),
                            showlegend=False,
                        )
                    )

                except (
                    TypeError,
                    ValueError,
                    KeyError,
                    IndexError,
                ):
                    continue

    # ==========================================================
    # 3. VẼ NODES
    # ==========================================================

    fig.add_trace(
        go.Scatter(
            x=nodes["x (m)"],
            y=nodes["y (m)"],
            mode="markers+text",
            text=[
                f"N{i}"
                for i in range(len(nodes))
            ],
            textposition="top center",
            marker=dict(
                size=10,
                color="#ffffff",
                line=dict(
                    color="#1f77b4",
                    width=2,
                ),
            ),
            name="Nodes",
            hovertemplate=(
                "Node: %{text}"
                "<br>X: %{x:.3f} m"
                "<br>Y: %{y:.3f} m"
                "<extra></extra>"
            ),
        )
    )

    # ==========================================================
    # 4. VẼ GỐI TỰA
    # ==========================================================

    if df_sup is not None and not df_sup.empty:

        supports = df_sup.copy()

        if "node" in supports.columns:

            for _, row in supports.iterrows():

                try:
                    node_id = int(row["node"])

                    if (
                        node_id < 0
                        or node_id >= len(nodes)
                    ):
                        continue

                    x = float(
                        nodes.loc[node_id, "x (m)"]
                    )

                    y = float(
                        nodes.loc[node_id, "y (m)"]
                    )

                    fig.add_trace(
                        go.Scatter(
                            x=[x],
                            y=[y],
                            mode="markers",
                            marker=dict(
                                symbol="triangle-up",
                                size=14,
                                color="#d62728",
                            ),
                            name="Support",
                            hovertemplate=(
                                f"Support at Node {node_id}"
                                "<extra></extra>"
                            ),
                            showlegend=False,
                        )
                    )

                except (
                    TypeError,
                    ValueError,
                    KeyError,
                    IndexError,
                ):
                    continue

    # ==========================================================
    # 5. ĐÁNH SỐ ELEMENT
    # ==========================================================

    if df_el is not None and not df_el.empty:

        elements = df_el.copy()

        if "i" in elements.columns and "j" in elements.columns:

            for e_id, row in elements.iterrows():

                try:
                    i = int(row["i"])
                    j = int(row["j"])

                    if (
                        i < 0
                        or j < 0
                        or i >= len(nodes)
                        or j >= len(nodes)
                    ):
                        continue

                    xi = float(nodes.loc[i, "x (m)"])
                    yi = float(nodes.loc[i, "y (m)"])

                    xj = float(nodes.loc[j, "x (m)"])
                    yj = float(nodes.loc[j, "y (m)"])

                    xm = (xi + xj) / 2
                    ym = (yi + yj) / 2

                    fig.add_annotation(
                        x=xm,
                        y=ym,
                        text=f"E{e_id}",
                        showarrow=False,
                        font=dict(
                            size=11,
                            color="#111111",
                        ),
                        bgcolor="rgba(255,255,255,0.7)",
                    )

                except (
                    TypeError,
                    ValueError,
                    KeyError,
                    IndexError,
                ):
                    continue

    # ==========================================================
    # 6. LAYOUT
    # ==========================================================

    fig.update_layout(
        title="Plane Frame Geometry",
        xaxis_title="X (m)",
        yaxis_title="Y (m)",
        template="plotly_white",
        height=500,
        margin=dict(
            l=50,
            r=30,
            t=60,
            b=50,
        ),
        xaxis=dict(
            zeroline=True,
            showgrid=True,
        ),
        yaxis=dict(
            zeroline=True,
            showgrid=True,
            scaleanchor="x",
            scaleratio=1,
        ),
        legend=dict(
            orientation="h",
            yanchor="bottom",
            y=1.02,
            xanchor="left",
            x=0,
        ),
    )

    return fig


# ══════════════════════════════════════════════════════
#  MAIN RUNNER
# ══════════════════════════════════════════════════════
def main():
    # Khởi tạo trạng thái sidebar lần đầu
    if "sidebar_open" not in st.session_state:
        st.session_state.sidebar_open = True

    inject_css()

    # ════════════════════════════════════════════════════════════
    # 🔥 ĐOẠN CSS VÁ LỖI: TRÁNH CHE KHUẤT LOGO & TIÊU ĐỀ TRÊN CÙNG 🔥
    # ════════════════════════════════════════════════════════════
    st.markdown("""
        <style>
        /* 1. Ẩn thanh header mặc định của Streamlit để không tranh chấp vị trí đè lên chữ */
        [data-testid="stHeader"] {
            display: none !important;
        }
        /* 2. Định nghĩa khoảng cách đỉnh đầu an toàn giúp tiêu đề cách mép trên một khoảng vừa vặn */
        .block-container {
            padding-top: 2rem !important;
        }
        </style>
    """, unsafe_allow_html=True)
    # ════════════════════════════════════════════════════════════

    # ── Header row: Nút ☰ nằm bên trái, Logo và Tiêu đề nằm ở giữa ──────────────────
    header_col1, header_col2, header_col3 = st.columns([0.05, 0.90, 0.05], vertical_alignment="center")

    with header_col1:
        # Nút taskbar (☰) nằm sát lề trái, thẳng hàng hoàn hảo với các tab bên dưới
        icon = "☰"
        if st.button(icon, key="toggle_sidebar", help="Mở/đóng bảng nhập liệu"):
            st.session_state.sidebar_open = not st.session_state.sidebar_open
            st.rerun()

    with header_col2:
        # Khối giữa chứa Logo và Tiêu đề được căn giữa tuyệt đối
        st.markdown(
            """
            <div style="
                display: flex;
                align-items: center;
                justify-content: center;
                gap: 25px;
                width: 100%;
            ">
            """,
            unsafe_allow_html=True
        )

        logo_sub_col, text_sub_col = st.columns([0.15, 0.85], vertical_alignment="center")

        with logo_sub_col:
            # Tăng kích thước logo to rõ (width=135)
            st.image("LOGO/TOOLBOX.png", width=135)

        with text_sub_col:
            # Tiêu đề to bản, sắc nét và chuyên nghiệp
            st.markdown(
                """
                <div style="
                    display: flex;
                    flex-direction: column;
                    justify-content: center;
                ">
                    <div style="
                        font-size: 32px;
                        font-weight: 800;
                        letter-spacing: 0.5px;
                        line-height: 1.2;
                        background: linear-gradient(90deg, #ffffff, #b0c4de);
                        -webkit-background-clip: text;
                        -webkit-text-fill-color: transparent;
                    ">
                        BEAM ANALYSIS TOOLBOX
                    </div>
                    <div style="
                        font-size: 15px;
                        opacity: 0.75;
                        margin-top: 5px;
                        font-weight: 400;
                        letter-spacing: 0.3px;
                    ">
                        Simplified Structural Analysis Tool
                    </div>
                </div>
                """,
                unsafe_allow_html=True
            )

        st.markdown("</div>", unsafe_allow_html=True)

    # Đường kẻ ngang phân tách header thanh lịch
    st.markdown("<hr style='margin-top: 10px; margin-bottom: 25px; border-color: rgba(255,255,255,0.1);'>",
                unsafe_allow_html=True)
    # ── Tabs ──────────────────────────────────────────
    # Thêm tab "Feedback" vào danh sách
    tab1, tab2, tab3, tab4 = st.tabs(["Single Beam", "Continuous Beam", "Plane Frame", "📋 Feedback"])
    with tab1: render_single_beam()
    with tab2: render_continuous_beam()
    with tab3: render_plane_frame()
    with tab4:
        st.title("📋 Đóng góp ý kiến & Báo lỗi")
        st.info("💡 Tôi luôn lắng nghe ý kiến từ bạn để cải tiến công cụ tốt hơn từng ngày!")

        st.markdown("""
        **Bạn có thể phản hồi về:**
        - 🐛 Các lỗi tính toán hoặc lỗi giao diện gặp phải.
        - 💡 Đề xuất thêm tính năng mới (Ví dụ: tính toán TTGH1, TTGH2,...).
        - 💬 Trải nghiệm tổng thể của bạn.
        """)


        st.link_button(
            "🚀 Gửi Feedback Qua Google Form",
            "https://forms.gle/fm8aSmCyX9LTPPsq9",
            type="primary",
            use_container_width=True
    )
if __name__ == "__main__":
    main()
